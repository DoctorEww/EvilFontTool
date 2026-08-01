import codecs
import copy
import logging
import os
import pathlib
import re
import string
import subprocess
import tempfile
import uuid
import zipfile
from xml.sax.saxutils import escape as _xml_escape

from docx import Document
from docx.document import Document as _DocxDocument
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from fontTools.ttLib import TTFont
from pdf2image import convert_from_path as _convert_from_path
from pdf2image.exceptions import PDFInfoNotInstalledError as _PDFInfoNotInstalledError
from pdfminer.high_level import extract_pages as _extract_pages
from pdfminer.layout import LTChar as _LTChar
from reportlab.lib.utils import ImageReader as _ImageReader
from reportlab.pdfbase import pdfmetrics as _pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont as _RLTTFont
from reportlab.pdfgen import canvas as _rl_canvas

logger = logging.getLogger(__name__)

_INK_FONT = "__ink__"
_REGION_TOL = 2.0



# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

# The set of characters to include in the font family.
# Modify this to limit which characters are processed.
LETTERS = (
    string.ascii_lowercase
    + " "
    + string.punctuation
    + string.ascii_uppercase
    + string.digits
)

# Advance width for invisible (stealth) characters.
# 0 works for most renderers, but some applications may behave unexpectedly.
WIDTH = 0

# Path to the bundled HTML template.
_TEMPLATE = pathlib.Path(__file__).parent / "data" / "template.html"


def _read_lines(path):
    """Read a user-supplied text file and split it into lines.

    A bare open(path) decodes with the platform-default encoding, which on
    Windows is not UTF-8 or UTF-16 -- and PowerShell's `>` / `echo` / `Out-File`
    (plus Notepad's "Unicode" option) write UTF-16 with a BOM. Reading that
    with the platform default doesn't raise; it silently produces a string
    full of embedded NUL bytes, which only surfaces later as a confusing lxml
    error. Sniff the BOM (if any) and decode accordingly instead.
    """
    with open(path, "rb") as handle:
        raw = handle.read()
    for bom, encoding in (
        (codecs.BOM_UTF8, "utf-8"),
        (codecs.BOM_UTF16_LE, "utf-16-le"),
        (codecs.BOM_UTF16_BE, "utf-16-be"),
    ):
        if raw.startswith(bom):
            # utf-16-le/-be don't strip their own BOM on decode (unlike
            # utf-8-sig for utf-8); slice it off ourselves in all three cases.
            return raw[len(bom):].decode(encoding).splitlines()
    return raw.decode("utf-8").splitlines()


# ---------------------------------------------------------------------------
# Font helpers
# ---------------------------------------------------------------------------

def _get_unicode_cmap(font):
    """Return the first suitable Windows Unicode cmap subtable (format 4 or 12).

    Raises ValueError if none is found.
    """
    for table in font['cmap'].tables:
        if (table.format in (4, 12)
                and table.platformID == 3
                and table.platEncID in (1, 10)):
            return table
    raise ValueError("No suitable Unicode cmap subtable found in the font.")


def _remove_layout_tables(font):
    """Remove OpenType tables that control ligatures, kerning, and substitution.

    These tables operate on glyph names rather than cmap entries, so they can
    override our remapping in unpredictable ways depending on surrounding characters.
    """
    for tag in ('GSUB', 'GPOS', 'kern', 'GDEF'):
        if tag in font:
            del font[tag]


def _rename_font(font, new_name):
    """Overwrite the family name records in the font's name table."""
    for record in font['name'].names:
        if record.nameID in (1, 4, 6):
            record.string = new_name.encode("utf-16-be")


# ---------------------------------------------------------------------------
# Core font generation
# ---------------------------------------------------------------------------

def createstealthfont(reference_font, output_dir, font_name):
    """Generate a stealth font where every character renders as an invisible space.

    The stealth font is used to hide the extra computer-file characters that
    have no corresponding human-file character to disguise themselves as.
    It is saved as both WOFF (web) and TTF (document/desktop).
    """
    font = TTFont(reference_font)
    unicode_cmap = _get_unicode_cmap(font)

    # Use the space glyph as the target — all other characters will map to it
    space_glyph_name = unicode_cmap.cmap.get(ord(" "))

    # Remap every character in LETTERS to the space glyph and zero its advance width
    for table in font['cmap'].tables:
        for char in LETTERS:
            if ord(char) in table.cmap:
                table.cmap[ord(char)] = space_glyph_name
                font['hmtx'].metrics[space_glyph_name] = (
                    WIDTH,
                    font['hmtx'].metrics[space_glyph_name][1],
                )

    _remove_layout_tables(font)

    # Internal font name for the stealth variant uses "0" as a sentinel
    new_name = f'{font_name} 0'
    _rename_font(font, new_name)
    logger.debug(f"  Stealth font internal name: {new_name}")

    # Append the @font-face CSS rule for the stealth font
    with open(f'{output_dir}/fonts.css', "a") as css_file:
        css_file.write(
            "@font-face {"
            'font-family: "0";'
            'src: url("fonts/0.woff") format(\'woff\');'
            "}"
        )

    # Save WOFF (flavor must be set explicitly) then TTF
    font.flavor = 'woff'
    font.save(f'{output_dir}/fonts/0.woff')
    font.flavor = None
    font.save(f'{output_dir}/ttffonts/0.ttf')

    logger.debug(f"[DONE] stealth font -> {output_dir}/fonts/0.woff + ttffonts/0.ttf")


def createfonts(reference_font, output_dir, font_name):
    """Generate one Evil Font variant per character in LETTERS.

    In each variant, every character's glyph is replaced with the glyph for
    `currentletter`. This means that no matter what Unicode byte is stored in
    the document, the renderer will draw `currentletter` — the core Evil Font trick.

    Also writes all @font-face rules to fonts.css (overwrites any existing file).
    """
    logger.debug(f"Source font:  {reference_font}")
    logger.debug(f"Output dir:   {output_dir}")
    logger.debug(f"Characters:   {len(LETTERS)} variants to generate")

    with open(f'{output_dir}/fonts.css', "w") as css_file:

        for currentletter in LETTERS:
            # Load a fresh copy of the font for each variant to avoid cross-contamination
            font = TTFont(reference_font)
            font.recalcBBoxes = False

            unicode_cmap = _get_unicode_cmap(font)

            # Get the source glyph and its advance width for this letter
            source_glyph_name = unicode_cmap.cmap.get(ord(currentletter))
            source_width = font['hmtx'].metrics[source_glyph_name][0]

            # Take a deep copy of the source glyph to use as a stamp
            source_glyph = copy.deepcopy(font['glyf'][source_glyph_name])

            # Remap every other character in LETTERS to look like currentletter
            for table in font['cmap'].tables:
                for char in LETTERS:
                    if char == currentletter:
                        continue
                    if ord(char) not in table.cmap:
                        continue

                    target_glyph_name = table.cmap[ord(char)]
                    target_original = font['glyf'][target_glyph_name]

                    # Preserve the target's original bounding box so that
                    # spacing and baseline positioning remain correct
                    orig_xMin = getattr(target_original, 'xMin', 0)
                    orig_yMax = getattr(target_original, 'yMax', 0)
                    orig_yMin = getattr(target_original, 'yMin', 0)

                    # Stamp a copy of the source glyph into the target slot
                    font['glyf'][target_glyph_name] = copy.deepcopy(source_glyph)

                    # Restore the vertical bounds (keeps line height consistent)
                    font['glyf'][target_glyph_name].yMax = orig_yMax
                    font['glyf'][target_glyph_name].yMin = orig_yMin

                    # Match advance width to source; preserve original LSB for positioning
                    font['hmtx'].metrics[target_glyph_name] = (source_width, orig_xMin)

            _remove_layout_tables(font)

            # Each font variant is named using the hex encoding of the letter
            # so the name is unique and safely usable as a filename
            letter_hex = currentletter.encode().hex()
            new_name = f'{font_name} {letter_hex}'
            _rename_font(font, new_name)

            # Save as WOFF for web use and TTF for document/desktop use
            font.flavor = 'woff'
            font.save(f'{output_dir}/fonts/{letter_hex}.woff')
            font.flavor = None
            font.save(f'{output_dir}/ttffonts/{letter_hex}.ttf')

            # Write the @font-face rule for this variant
            css_file.write(
                f'@font-face {{'
                f'font-family: "{letter_hex}";'
                f'src: url("fonts/{letter_hex}.woff") format(\'woff\');'
                f'}}'
            )

            logger.debug(f"  [{letter_hex}] '{currentletter}' -> {output_dir}/fonts/{letter_hex}.woff + ttffonts/{letter_hex}.ttf")


# ---------------------------------------------------------------------------
# HTML output
# ---------------------------------------------------------------------------

def _write_html(output_file, content):
    """Inject `content` into the bundled HTML template and write to `output_file`."""
    html = _TEMPLATE.read_text()
    html = html.replace("<!-- #STUFF HERE -->", content)
    with open(output_file, "w") as f:
        f.write(html)


def createhtml(input_human_file, input_computer_file, output_file):
    """Build a steganographic HTML file from human and computer text files.

    Each character from the computer file is wrapped in a <span> that applies
    an Evil Font chosen by the corresponding human file character. To a human
    reading the rendered page, the text looks like the human file. A machine
    parsing the raw HTML sees the computer file.

    Lines where the computer text is shorter than the human text are skipped.
    Extra computer characters (beyond the human line length) are hidden using
    the stealth font (font-family: '0').
    """
    logger.debug(f"Human file:    {input_human_file}")
    logger.debug(f"Computer file: {input_computer_file}")
    logger.debug(f"Output file:   {output_file}")

    spans = []
    lines_processed = 0
    total_hidden = 0
    line_number = 0

    for human_line, computer_line in zip(
        _read_lines(input_human_file),
        _read_lines(input_computer_file),
    ):
        h_len = len(human_line)
        c_len = len(computer_line)

        line_number += 1
        logger.debug(f"Human   ({h_len} chars): {human_line}")
        logger.debug(f"Computer({c_len} chars): {computer_line}")

        if c_len < h_len:
            logger.warning(f"  Skipping line {line_number}: computer line must be >= human line length.")
            continue

        # Extra computer characters are inserted at the midpoint of the human text
        diff = c_len - h_len
        mid = h_len // 2

        logger.debug(f"  diff={diff}, hidden chars inserted at mid={mid}")

        line_spans = []
        h_index = 0

        for i, computer_char in enumerate(computer_line):
            if mid <= i < mid + diff:
                # Hidden character — rendered invisibly using the stealth font
                line_spans.append(
                    f"<span style=\"font-family: '0';\">{computer_char}</span>"
                )
            else:
                # Visible character — disguised as the corresponding human character
                human_char = human_line[h_index]
                letter_hex = human_char.encode().hex()
                line_spans.append(
                    f"<span style=\"font-family: '{letter_hex}';\">{computer_char}</span>"
                )
                h_index += 1

        spans.append("".join(line_spans))
        lines_processed += 1
        total_hidden += diff

    _write_html(output_file, "\n<br>\n".join(spans))
    logger.debug(f"[DONE] HTML written -> {output_file} ({lines_processed} lines, {total_hidden} hidden chars)")


# ---------------------------------------------------------------------------
# DOCX font embedding
# ---------------------------------------------------------------------------

# OOXML content type for an obfuscated embedded font part (ECMA-376 Part 1, 17.9).
_FONTDATA_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.obfuscatedFont"


def _obfuscate_font_bytes(data, guid_bytes):
    """XOR the first 32 bytes of TTF `data` with `guid_bytes`, taken in reverse.

    This is the OOXML font-embedding obfuscation algorithm (ECMA-376 Part 1,
    17.9.2). Word reverses it with the identical XOR -- using the GUID stored
    as the embedded part's fontKey -- before treating the bytes as a normal
    TTF; XOR being its own inverse is exactly why one function does both
    directions.
    """
    data = bytearray(data)
    for i in range(min(32, len(data))):
        data[i] ^= guid_bytes[15 - (i % 16)]
    return bytes(data)


def _next_rel_id(rels_xml):
    """Return an rIdN not already used in a .rels part's raw XML text."""
    used = {int(n) for n in re.findall(r'Id="rId(\d+)"', rels_xml)}
    n = 1
    while n in used:
        n += 1
    return f"rId{n}"


def _embed_fonts(docx_path, ttf_dir, font_names):
    """Embed the given Evil Font TTFs directly into a saved .docx, in place.

    `font_names` are the exact font family names used in the document's runs
    (e.g. 'MyFont 68', 'MyFont 0'); each one's TTF is found in `ttf_dir` by the
    hex/'0' suffix after the last space -- the same naming createfonts() and
    createstealthfont() use. This builds the OOXML font-embedding parts
    (fontTable.xml entries, obfuscated font data, relationships, content
    types) directly, rather than relying on Word's -- or LibreOffice's, which
    doesn't work for Evil Fonts -- own "embed fonts" save option.

    python-docx's default template always ships a word/fontTable.xml (listing
    the built-in style fonts) and a document.xml.rels entry for it already, so
    only new <w:font> entries need appending; nothing else needs to reference
    fontTable.xml for the first time.
    """
    with zipfile.ZipFile(docx_path, "r") as zin:
        parts = {name: zin.read(name) for name in zin.namelist()}

    fonttable_xml = parts["word/fontTable.xml"].decode("utf-8")
    fonttable_rels_xml = parts.get(
        "word/_rels/fontTable.xml.rels",
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>',
    ).decode("utf-8")
    content_types_xml = parts["[Content_Types].xml"].decode("utf-8")

    new_fonts, new_rels, new_overrides = [], [], []
    idx = 0
    for font_name in sorted(font_names):
        suffix = font_name.rsplit(" ", 1)[-1]
        ttf_path = os.path.join(ttf_dir, f"{suffix}.ttf")
        if not os.path.isfile(ttf_path):
            logger.warning("embed: no TTF for font '%s' at %s -- skipping.", font_name, ttf_path)
            continue

        idx += 1
        part_name = f"fonts/font{idx}.fntdata"
        rid = _next_rel_id(fonttable_rels_xml + "".join(new_rels))

        guid = uuid.uuid4()
        with open(ttf_path, "rb") as handle:
            raw = handle.read()
        parts[f"word/{part_name}"] = _obfuscate_font_bytes(raw, guid.bytes)

        new_fonts.append(
            f'<w:font w:name="{_xml_escape(font_name)}">'
            f'<w:embedRegular r:id="{rid}" w:fontKey="{{{str(guid).upper()}}}" w:subsetted="false"/>'
            f'</w:font>'
        )
        new_rels.append(
            f'<Relationship Id="{rid}" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" '
            f'Target="{part_name}"/>'
        )
        new_overrides.append(
            f'<Override PartName="/word/{part_name}" ContentType="{_FONTDATA_CONTENT_TYPE}"/>'
        )

    if not new_fonts:
        logger.warning("embed: no fonts embedded (none of the used fonts had a matching TTF).")
        return

    parts["word/fontTable.xml"] = fonttable_xml.replace(
        "</w:fonts>", "".join(new_fonts) + "</w:fonts>").encode("utf-8")
    parts["word/_rels/fontTable.xml.rels"] = fonttable_rels_xml.replace(
        "</Relationships>", "".join(new_rels) + "</Relationships>").encode("utf-8")
    parts["[Content_Types].xml"] = content_types_xml.replace(
        "</Types>", "".join(new_overrides) + "</Types>").encode("utf-8")

    settings_xml = parts["word/settings.xml"].decode("utf-8")
    if "<w:embedTrueTypeFonts" not in settings_xml:
        if "<w:proofState" in settings_xml:
            settings_xml = settings_xml.replace(
                "<w:proofState", "<w:embedTrueTypeFonts/><w:proofState", 1)
        else:
            insert_at = settings_xml.index(">", settings_xml.index("<w:settings")) + 1
            settings_xml = (settings_xml[:insert_at] + "<w:embedTrueTypeFonts/>"
                            + settings_xml[insert_at:])
        parts["word/settings.xml"] = settings_xml.encode("utf-8")

    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)

    logger.debug(f"[DONE] Embedded {idx} font(s) into {docx_path}")


# ---------------------------------------------------------------------------
# DOCX output
# ---------------------------------------------------------------------------

def create_doc(input_human_file, input_computer_file, output_file, font_name_in,
               author="anonymous", ttf_dir=None):
    """Build a steganographic DOCX file from human and computer text files.

    Works identically to createhtml() but outputs a Word document. Each run is
    assigned an Evil Font by setting all four font slots (ascii, hAnsi, eastAsia,
    cs) to prevent Word from falling back to a system font.

    If `ttf_dir` is given (the `ttffonts` directory from the `create` step), the
    Evil Font TTFs actually used are embedded directly into the saved .docx, so
    the deception renders correctly even where the fonts aren't installed.
    Otherwise the TTF variants must be installed on the system for the
    deception to render correctly in Word.

    Lines where the computer text is shorter than the human text are skipped.
    """
    logger.debug(f"Human file:    {input_human_file}")
    logger.debug(f"Computer file: {input_computer_file}")
    logger.debug(f"Output file:   {output_file}")
    logger.debug(f"Font family:   {font_name_in}")

    doc = Document()
    lines_processed = 0
    total_hidden = 0
    line_number = 0
    used_font_names = set()

    for human_line, computer_line in zip(
        _read_lines(input_human_file),
        _read_lines(input_computer_file),
    ):
        h_len = len(human_line)
        c_len = len(computer_line)

        line_number += 1
        logger.debug(f"Human   ({h_len} chars): {human_line}")
        logger.debug(f"Computer({c_len} chars): {computer_line}")

        if c_len < h_len:
            logger.warning(f"  Skipping line {line_number}: computer line must be >= human line length.")
            continue

        diff = c_len - h_len
        mid = h_len // 2

        logger.debug(f"  diff={diff}, hidden chars inserted at mid={mid}")

        p = doc.add_paragraph()
        h_index = 0

        for i, computer_char in enumerate(computer_line):
            if mid <= i < mid + diff:
                # Hidden character — use the stealth (zero-width) font
                font_name = f'{font_name_in} 0'
            else:
                # Visible character — disguised as the corresponding human character
                human_char = human_line[h_index]
                font_name = f'{font_name_in} {human_char.encode().hex()}'
                h_index += 1

            # Add a run and explicitly set all four font slots.
            # Word will fall back to a system font if any slot is unset,
            # which would break the illusion.
            run = p.add_run(computer_char)
            run.font.name = font_name
            rFonts = run._element.rPr.rFonts
            rFonts.set(qn("w:ascii"),   font_name)
            rFonts.set(qn("w:hAnsi"),   font_name)
            rFonts.set(qn("w:eastAsia"), font_name)
            rFonts.set(qn("w:cs"),      font_name)
            used_font_names.add(font_name)

        lines_processed += 1
        total_hidden += diff

    doc.core_properties.comments = ""
    if author is not None:
        doc.core_properties.author = author

    doc.save(output_file)

    if ttf_dir:
        _embed_fonts(output_file, ttf_dir, used_font_names)

    logger.debug(f"[DONE] DOCX written -> {output_file} ({lines_processed} lines, {total_hidden} hidden chars)")

# ---------------------------------------------------------------------------
# PDF output (copy-paste-safe two-layer render)
# ---------------------------------------------------------------------------
#
# Turns an Evil Font DOCX (produced by create_doc) into a PDF that:
#   * looks EXACTLY like the docx  -- LibreOffice renders the real document,
#     so all formatting/headers/footers are preserved and rasterised as a
#     non-selectable image layer, and
#   * copy-pastes the hidden payload correctly in every viewer (incl. poppler),
#     by squishing each rendered line's payload (visible + hidden chars) onto
#     that line's exact box as invisible text (PDF render mode 3).
#
# Runtime requirements (install once):
#   * LibreOffice ('soffice' on PATH) and poppler-utils
#   * pip install reportlab pdf2image pdfminer.six Pillow




# ============================================================================
# DOCX PARSING
# ============================================================================

def _iter_block_items(parent):
    """
    Yield every paragraph in the document in reading order.

    Includes paragraphs inside tables, including nested tables.
    """

    elem = (
        parent.element.body
        if isinstance(parent, _DocxDocument)
        else parent._tc
    )

    for child in elem.iterchildren():

        if isinstance(child, CT_P):

            yield Paragraph(child, parent)

        elif isinstance(child, CT_Tbl):

            table = Table(child, parent)

            for row in table.rows:

                for cell in row.cells:

                    yield from _iter_block_items(cell)


def _run_font_name(run):
    """
    Return the font family assigned to a run.
    """

    name = run.font.name

    if name:
        return name

    rpr = run._element.rPr

    if rpr is not None:

        rfonts = rpr.find(
            qn("w:rFonts")
        )

        if rfonts is not None:

            return (
                rfonts.get(qn("w:ascii"))
                or rfonts.get(qn("w:hAnsi"))
                or ""
            )

    return ""


def _is_stealth_font(font_name):
    """
    Detect the stealth font.

    Stealth fonts are named with a suffix of:

        " 0"

    Example:

        EvilFont 0
    """

    base = (
        font_name or ""
    ).split(
        "+",
        1
    )[-1]

    return base.endswith(
        " 0"
    )


def _is_decorative_glyph(char_text):
    """
    Detect native decorative glyphs such as list bullets.
    """

    return (
        len(char_text) == 1
        and 0xE000 <= ord(char_text) <= 0xF8FF
    )


def _paragraph_payload(paragraphs):
    """
    Flatten paragraphs into:

        payload:
            [(character, hidden), ...]

        visible:
            all visible DOCX characters

    Hidden Evil Font characters are preserved in payload but are not
    added to visible because they should not consume a rendered PDF glyph.
    """

    payload = []
    visible = []

    for paragraph in paragraphs:

        for run in paragraph.runs:

            hidden = _is_stealth_font(
                _run_font_name(run)
            )

            for char in run.text:

                payload.append(
                    (
                        char,
                        hidden
                    )
                )

                if not hidden:

                    visible.append(
                        char
                    )

    return payload, "".join(visible)


def _parse_docx_payload(docx_path):
    """
    Extract body/header/footer payloads and page geometry.
    """

    doc = Document(
        docx_path
    )

    body = _paragraph_payload(
        list(
            _iter_block_items(doc)
        )
    )

    section = doc.sections[0]

    header = _paragraph_payload(
        section.header.paragraphs
    )

    footer = _paragraph_payload(
        section.footer.paragraphs
    )

    top = (
        section.top_margin.pt
        if section.top_margin is not None
        else 72.0
    )

    bottom = (
        section.bottom_margin.pt
        if section.bottom_margin is not None
        else 72.0
    )

    return {
        "body": body,
        "header": header,
        "footer": footer,
        "top_margin": top,
        "bottom_margin": bottom,
    }


# ============================================================================
# LIBREOFFICE RENDERING
# ============================================================================

def _render_docx_to_pdf(
    docx_path,
    ttf_dir,
    workdir,
    soffice
):
    """
    Render the original DOCX through LibreOffice.
    """

    docx_path = os.path.abspath(
        docx_path
    )

    env = dict(
        os.environ
    )

    if ttf_dir:

        conf = os.path.join(
            workdir,
            "fonts.conf"
        )

        cache = os.path.join(
            workdir,
            "fccache"
        )

        os.makedirs(
            cache,
            exist_ok=True
        )

        with open(
            conf,
            "w"
        ) as handle:

            handle.write(
                '<?xml version="1.0"?>\n'
                '<!DOCTYPE fontconfig SYSTEM "fonts.dtd">\n'
                "<fontconfig>\n"
                f"  <dir>{os.path.abspath(ttf_dir)}</dir>\n"
                f"  <cachedir>{cache}</cachedir>\n"
                '  <include ignore_missing="yes">'
                '/etc/fonts/fonts.conf'
                '</include>\n'
                "</fontconfig>\n"
            )

        env[
            "FONTCONFIG_FILE"
        ] = conf

        try:

            subprocess.run(
                [
                    "fc-cache",
                    "-f"
                ],
                env=env,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL
            )

        except FileNotFoundError:

            logger.warning(
                "fc-cache not found. "
                "Skipping font cache refresh."
            )

    profile = (
        pathlib.Path(
            workdir,
            "loprofile"
        )
        .resolve()
        .as_uri()
    )

    soffice_dir = (
        os.path.dirname(
            soffice
        )
        or None
    )

    try:

        result = subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                workdir,
                docx_path,
                f"-env:UserInstallation={profile}",
            ],
            env=env,
            cwd=soffice_dir,
            capture_output=True,
            text=True,
        )

    except FileNotFoundError:

        raise RuntimeError(
            f"Could not find LibreOffice binary: {soffice}"
        )

    if result.returncode != 0:

        detail = (
            result.stderr
            or result.stdout
            or ""
        ).strip()

        raise RuntimeError(
            "LibreOffice failed to convert DOCX."
            + (
                f"\n{detail}"
                if detail
                else ""
            )
        )

    output_pdf = os.path.join(
        workdir,
        os.path.splitext(
            os.path.basename(
                docx_path
            )
        )[0]
        + ".pdf"
    )

    if not os.path.exists(
        output_pdf
    ):

        detail = (
            result.stderr
            or result.stdout
            or ""
        ).strip()

        raise RuntimeError(
            "LibreOffice produced no PDF."
            + (
                f"\n{detail}"
                if detail
                else ""
            )
        )

    return output_pdf


# ============================================================================
# PDF GLYPH EXTRACTION
# ============================================================================

def _split_row_by_gaps(
    row,
    gap_factor=2.5
):
    """
    Split a row into segments when large gaps occur.
    """

    if not row:

        return []

    segments = []

    current = [
        row[0]
    ]

    for prev, glyph in zip(
        row,
        row[1:]
    ):

        gap = (
            glyph.x0
            - prev.x1
        )

        if gap > gap_factor * max(
            prev.size,
            glyph.size,
            1.0
        ):

            segments.append(
                current
            )

            current = [
                glyph
            ]

        else:

            current.append(
                glyph
            )

    segments.append(
        current
    )

    return segments


def _visible_lines(
    pdf_path,
    y_tol=3.0
):
    """
    Recover rendered visible PDF glyphs.

    Each line contains:

        box:
            bounding box

        n:
            number of visible glyphs

        chars:
            actual rendered PDF characters
    """

    pages = []

    for layout in _extract_pages(
        pdf_path
    ):

        glyphs = []

        stack = [
            layout
        ]

        while stack:

            obj = stack.pop()

            if isinstance(
                obj,
                _LTChar
            ):

                glyphs.append(
                    obj
                )

            elif hasattr(
                obj,
                "__iter__"
            ):

                try:

                    stack.extend(
                        list(obj)
                    )

                except TypeError:

                    pass

        # ---------------------------------------------------------------
        # Filter visible glyphs
        # ---------------------------------------------------------------

        visible = []

        for g in glyphs:

            if _is_stealth_font(
                g.fontname
            ):

                continue

            if (
                g.x1
                - g.x0
            ) <= 0.1 * max(
                g.size,
                1.0
            ):

                continue

            if _is_decorative_glyph(
                g.get_text()
            ):

                continue

            visible.append(
                g
            )

        # ---------------------------------------------------------------
        # Sort top-to-bottom and left-to-right
        # ---------------------------------------------------------------

        visible.sort(
            key=lambda g: (
                -g.y0,
                g.x0
            )
        )

        rows = []

        current = []

        base_y = None

        for g in visible:

            if (
                base_y is None
                or abs(
                    g.y0
                    - base_y
                ) <= y_tol
            ):

                current.append(
                    g
                )

                if base_y is None:

                    base_y = g.y0

            else:

                rows.append(
                    current
                )

                current = [
                    g
                ]

                base_y = g.y0

        if current:

            rows.append(
                current
            )

        lines = []

        for row in rows:

            row.sort(
                key=lambda g: g.x0
            )

            segments = _split_row_by_gaps(
                row
            )

            for segment in segments:

                box = (
                    min(
                        g.x0
                        for g in segment
                    ),

                    min(
                        g.y0
                        for g in segment
                    ),

                    max(
                        g.x1
                        for g in segment
                    ),

                    max(
                        g.y1
                        for g in segment
                    )
                )

                chars = [
                    g.get_text()
                    for g in segment
                ]

                lines.append(
                    {
                        "box": box,
                        "n": len(chars),
                        "chars": chars,
                    }
                )

        pages.append(
            {
                "size": (
                    layout.width,
                    layout.height
                ),
                "lines": lines,
            }
        )

    return pages


# ============================================================================
# CHARACTER MATCHING
# ============================================================================

def _chars_match(
    docx_char,
    pdf_char
):
    """
    Determine whether a DOCX character and PDF glyph represent
    the same visible character.
    """

    return (
        docx_char
        == pdf_char
    )


def _assign_payload(
    payload,
    lines,
    debug=False
):
    """
    Assign DOCX payload to rendered PDF lines one character at a time.

    Debug output only reports alignment problems and recovery actions.
    Successful character matches are silent.
    """

    buffers = [
        ""
        for _ in lines
    ]

    # ---------------------------------------------------------------
    # Flatten rendered PDF glyphs
    # ---------------------------------------------------------------

    rendered = []

    for line_index, line in enumerate(
        lines
    ):

        for glyph_index, char in enumerate(
            line["chars"]
        ):

            rendered.append(
                {
                    "line_index": line_index,
                    "glyph_index": glyph_index,
                    "char": char,
                }
            )

    pdf_index = 0

    mismatch_count = 0
    pdf_extra_count = 0
    missing_glyph_count = 0
    exhausted_count = 0

    # ---------------------------------------------------------------
    # Walk DOCX payload character-by-character
    # ---------------------------------------------------------------

    for payload_index, (
        docx_char,
        hidden
    ) in enumerate(
        payload
    ):

        # -----------------------------------------------------------
        # Hidden Evil Font character
        # -----------------------------------------------------------

        if hidden:

            if pdf_index > 0:

                line_index = rendered[
                    pdf_index - 1
                ][
                    "line_index"
                ]

            elif rendered:

                line_index = rendered[
                    0
                ][
                    "line_index"
                ]

            else:

                line_index = (
                    len(lines)
                    - 1
                )

            if line_index >= 0:

                buffers[
                    line_index
                ] += docx_char

            continue

        # -----------------------------------------------------------
        # PDF has no more glyphs
        # -----------------------------------------------------------

        if pdf_index >= len(
            rendered
        ):

            exhausted_count += 1

            if buffers:

                buffers[
                    -1
                ] += docx_char

            if debug:

                print(
                    "[PDF EXHAUSTED] "
                    f"DOCX={repr(docx_char)} "
                    f"payload_index={payload_index}"
                )

            continue

        # -----------------------------------------------------------
        # Current PDF glyph
        # -----------------------------------------------------------

        pdf_glyph = rendered[
            pdf_index
        ]

        pdf_char = pdf_glyph[
            "char"
        ]

        line_index = pdf_glyph[
            "line_index"
        ]

        # -----------------------------------------------------------
        # Normal match
        # -----------------------------------------------------------

        if _chars_match(
            docx_char,
            pdf_char
        ):

            buffers[
                line_index
            ] += docx_char

            pdf_index += 1

            continue

        # -----------------------------------------------------------
        # Mismatch
        # -----------------------------------------------------------

        mismatch_count += 1

        if debug:

            print(
                "[MISMATCH] "
                f"DOCX={repr(docx_char)} "
                f"PDF={repr(pdf_char)} "
                f"payload_index={payload_index} "
                f"pdf_index={pdf_index}"
            )

        # -----------------------------------------------------------
        # Look ahead for DOCX character in PDF
        # -----------------------------------------------------------

        found_index = None

        for lookahead in range(
            1,
            6
        ):

            candidate_index = (
                pdf_index
                + lookahead
            )

            if (
                candidate_index
                >= len(rendered)
            ):

                break

            candidate_char = rendered[
                candidate_index
            ][
                "char"
            ]

            if _chars_match(
                docx_char,
                candidate_char
            ):

                found_index = (
                    candidate_index
                )

                break

        # -----------------------------------------------------------
        # PDF has extra glyphs
        # -----------------------------------------------------------

        if found_index is not None:

            skipped = "".join(
                item[
                    "char"
                ]
                for item in rendered[
                    pdf_index:found_index
                ]
            )

            pdf_extra_count += len(
                skipped
            )

            if debug:

                print(
                    "[PDF EXTRA GLYPHS] "
                    f"skipping={repr(skipped)} "
                    f"before={repr(docx_char)}"
                )

            pdf_index = found_index

            line_index = rendered[
                pdf_index
            ][
                "line_index"
            ]

            buffers[
                line_index
            ] += docx_char

            pdf_index += 1

            continue

        # -----------------------------------------------------------
        # DOCX character has no matching PDF glyph
        # -----------------------------------------------------------

        missing_glyph_count += 1

        if debug:

            print(
                "[DOCX GLYPH MISSING] "
                f"preserving={repr(docx_char)} "
                f"against_pdf={repr(pdf_char)}"
            )

        buffers[
            line_index
        ] += docx_char

        # Do not advance PDF index.
        #
        # The next DOCX character will continue trying to match
        # the current PDF glyph.

    # ---------------------------------------------------------------
    # Summary
    # ---------------------------------------------------------------

    if debug:

        print(
            "\n"
            "[ALIGNMENT SUMMARY] "
            f"mismatches={mismatch_count}, "
            f"pdf_extra_glyphs={pdf_extra_count}, "
            f"docx_missing_glyphs={missing_glyph_count}, "
            f"pdf_exhausted_chars={exhausted_count}"
        )

    return buffers


# ============================================================================
# INVISIBLE PDF TEXT
# ============================================================================

def _draw_invisible(
    pdf,
    text,
    box,
    min_size=4.0,
    max_size=144.0
):
    """
    Draw invisible but selectable text over a line.
    """

    if not text:

        return

    x0, y0, x1, y1 = box

    width = max(
        1.0,
        x1 - x0
    )

    size = max(
        min_size,
        min(
            max_size,
            y1 - y0
        )
    )

    natural = _pdfmetrics.stringWidth(
        text,
        _INK_FONT,
        size
    )

    if not natural:

        natural = 1.0

    text_obj = pdf.beginText(
        x0,
        y0
    )

    text_obj.setFont(
        _INK_FONT,
        size
    )

    text_obj.setTextRenderMode(
        3
    )

    text_obj.setHorizScale(
        width
        / natural
        * 100.0
    )

    text_obj.textOut(
        text
    )

    pdf.drawText(
        text_obj
    )


# ============================================================================
# INK FONT RESOLUTION
# ============================================================================

def _resolve_ink_font(
    ink_font
):
    """
    Find a Unicode-capable font for invisible PDF text.
    """

    if ink_font:

        return ink_font

    for query in (
        "DejaVu Sans",
        "Liberation Sans",
        "sans-serif",
    ):

        try:

            found = subprocess.run(
                [
                    "fc-match",
                    "-f",
                    "%{file}",
                    query,
                ],
                capture_output=True,
                text=True,
            ).stdout.strip()

            if (
                found
                and os.path.exists(
                    found
                )
            ):

                return found

        except FileNotFoundError:

            break

    raise RuntimeError(
        "Could not find a Unicode-capable ink font. "
        "Pass --ink-font explicitly."
    )


# ============================================================================
# MAIN PDF CREATION
# ============================================================================

def create_pdf(
    input_docx,
    output_pdf,
    ttf_dir=None,
    dpi=200,
    soffice="soffice",
    ink_font=None,
    title="Untitled",
    author=None,
    subject=None,
    producer=None,
    debug=False
):
    """
    Convert Evil Font DOCX to PDF.

    Visible appearance:
        Rasterized image of the LibreOffice-rendered DOCX.

    Copyable text:
        Invisible PDF text layer containing the complete DOCX payload,
        including hidden Evil Font characters.
    """

    # ---------------------------------------------------------------
    # Register invisible text font
    # ---------------------------------------------------------------

    _pdfmetrics.registerFont(
        _RLTTFont(
            _INK_FONT,
            _resolve_ink_font(
                ink_font
            )
        )
    )

    # ---------------------------------------------------------------
    # Parse DOCX payload
    # ---------------------------------------------------------------

    info = _parse_docx_payload(
        input_docx
    )

    body_payload, body_visible = info[
        "body"
    ]

    header_payload, header_visible = info[
        "header"
    ]

    footer_payload, footer_visible = info[
        "footer"
    ]

    top_margin = info[
        "top_margin"
    ]

    bottom_margin = info[
        "bottom_margin"
    ]

    has_header = bool(
        header_payload
    )

    has_footer = bool(
        footer_payload
    )

    # ---------------------------------------------------------------
    # Render DOCX through LibreOffice
    # ---------------------------------------------------------------

    with tempfile.TemporaryDirectory() as workdir:

        look_pdf = _render_docx_to_pdf(
            input_docx,
            ttf_dir,
            workdir,
            soffice
        )

        pages = _visible_lines(
            look_pdf
        )

        try:

            images = _convert_from_path(
                look_pdf,
                dpi=dpi
            )

        except _PDFInfoNotInstalledError:

            raise RuntimeError(
                "Poppler is not installed or not available on PATH."
            )

    # ---------------------------------------------------------------
    # Classify rendered lines
    # ---------------------------------------------------------------

    page_regions = []

    for page in pages:

        page_height = page[
            "size"
        ][1]

        regions = {
            "header": [],
            "body": [],
            "footer": [],
        }

        for line in page[
            "lines"
        ]:

            y0 = line[
                "box"
            ][1]

            if (
                has_header
                and y0
                >= page_height
                - top_margin
                - _REGION_TOL
            ):

                regions[
                    "header"
                ].append(
                    line
                )

            elif (
                has_footer
                and y0
                <= bottom_margin
                + _REGION_TOL
            ):

                regions[
                    "footer"
                ].append(
                    line
                )

            else:

                regions[
                    "body"
                ].append(
                    line
                )

        page_regions.append(
            regions
        )

    # ---------------------------------------------------------------
    # Match body character-by-character
    # ---------------------------------------------------------------

    body_lines = [
        line
        for regions in page_regions
        for line in regions[
            "body"
        ]
    ]

    body_text = _assign_payload(
        body_payload,
        body_lines,
        debug=debug
    )

    # ---------------------------------------------------------------
    # Create output PDF
    # ---------------------------------------------------------------

    pdf = _rl_canvas.Canvas(
        output_pdf
    )

    if title:

        pdf.setTitle(
            title
        )

    if author:

        pdf.setAuthor(
            author
        )

    if subject is not None:

        pdf.setSubject(
            subject
        )

    if producer is not None:

        pdf.setProducer(
            producer
        )

    body_index = 0

    for page_number, page in enumerate(
        pages
    ):

        page_width, page_height = page[
            "size"
        ]

        pdf.setPageSize(
            (
                page_width,
                page_height
            )
        )

        # -----------------------------------------------------------
        # Visible rasterized page
        # -----------------------------------------------------------

        pdf.drawImage(
            _ImageReader(
                images[
                    page_number
                ]
            ),
            0,
            0,
            width=page_width,
            height=page_height,
        )

        regions = page_regions[
            page_number
        ]

        # -----------------------------------------------------------
        # Header/footer payload
        # -----------------------------------------------------------

        for region_payload, region_lines in (

            (
                header_payload,
                regions[
                    "header"
                ]
            ),

            (
                footer_payload,
                regions[
                    "footer"
                ]
            ),

        ):

            region_text = _assign_payload(
                region_payload,
                region_lines,
                debug=debug
            )

            for line, text in zip(
                region_lines,
                region_text
            ):

                _draw_invisible(
                    pdf,
                    text,
                    line[
                        "box"
                    ]
                )

        # -----------------------------------------------------------
        # Body payload
        # -----------------------------------------------------------

        for line in regions[
            "body"
        ]:

            if body_index < len(
                body_text
            ):

                _draw_invisible(
                    pdf,
                    body_text[
                        body_index
                    ],
                    line[
                        "box"
                    ]
                )

            body_index += 1

        pdf.showPage()

    pdf.save()

    logger.debug(
        "[DONE] PDF written -> %s",
        output_pdf
    )

    return output_pdf
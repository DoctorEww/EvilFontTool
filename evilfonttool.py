import argparse
import sys
import os
import copy
import string

from fontTools.ttLib import TTFont
from docx import Document
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Verbosity
# ---------------------------------------------------------------------------

# Set to True by passing -v / --verbose on the command line.
# All informational output is gated through log() so silent mode is clean.
VERBOSE = False


def log(msg, verbose_only=True):
    """Print msg to stdout.

    If verbose_only is True (default), only prints when VERBOSE is enabled.
    Pass verbose_only=False for messages that should always appear regardless
    of verbosity — e.g. skipped-line warnings that affect output correctness.
    Errors in main() always use print() directly and are never silenced.
    """
    if not verbose_only or VERBOSE:
        print(msg)


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

# The set of characters to include in the font family.
# Modify this to limit which characters are processed.
LETTERS = (
    string.ascii_lowercase
    + " "
    + string.punctuation
    + string.ascii_uppercase
    + string.digits
)

# Advance width for invisible (stealth) characters.
# 0 works for most renderers, but some applications may behave unexpectedly.
WIDTH = 0


# ---------------------------------------------------------------------------
# Font helpers
# ---------------------------------------------------------------------------

def _get_unicode_cmap(font):
    """Return the first suitable Windows Unicode cmap subtable (format 4 or 12).
    
    Raises ValueError if none is found.
    """
    for table in font['cmap'].tables:
        if (table.format in (4, 12)
                and table.platformID == 3
                and table.platEncID in (1, 10)):
            return table
    raise ValueError("No suitable Unicode cmap subtable found in the font.")


def _remove_layout_tables(font):
    """Remove OpenType tables that control ligatures, kerning, and substitution.
    
    These tables operate on glyph names rather than cmap entries, so they can
    override our remapping in unpredictable ways depending on surrounding characters.
    """
    for tag in ('GSUB', 'GPOS', 'kern', 'GDEF'):
        if tag in font:
            del font[tag]


def _rename_font(font, new_name):
    """Overwrite the family name records in the font's name table."""
    for record in font['name'].names:
        if record.nameID in (1, 4, 6):
            record.string = new_name.encode("utf-16-be")


# ---------------------------------------------------------------------------
# Core font generation
# ---------------------------------------------------------------------------

def createstealthfont(reference_font, output_dir, font_name):
    """Generate a stealth font where every character renders as an invisible space.
    
    The stealth font is used to hide the extra computer-file characters that
    have no corresponding human-file character to disguise themselves as.
    It is saved as both WOFF (web) and TTF (document/desktop).
    """
    font = TTFont(reference_font)
    unicode_cmap = _get_unicode_cmap(font)

    # Use the space glyph as the target — all other characters will map to it
    space_glyph_name = unicode_cmap.cmap.get(ord(" "))

    # Remap every character in LETTERS to the space glyph and zero its advance width
    for table in font['cmap'].tables:
        for char in LETTERS:
            if ord(char) in table.cmap:
                table.cmap[ord(char)] = space_glyph_name
                font['hmtx'].metrics[space_glyph_name] = (
                    WIDTH,
                    font['hmtx'].metrics[space_glyph_name][1],
                )

    _remove_layout_tables(font)

    # Internal font name for the stealth variant uses "0" as a sentinel
    new_name = f'{font_name} 0'
    _rename_font(font, new_name)
    log(f"  Stealth font internal name: {new_name}")

    # Append the @font-face CSS rule for the stealth font
    with open(f'{output_dir}/fonts.css', "a") as css_file:
        css_file.write(
            "@font-face {"
            'font-family: "0";'
            'src: url("fonts/0.woff") format(\'woff\');'
            "}"
        )

    # Save WOFF (flavor must be set explicitly) then TTF
    font.flavor = 'woff'
    font.save(f'{output_dir}/fonts/0.woff')
    font.flavor = None
    font.save(f'{output_dir}/ttffonts/0.ttf')

    log(f"[DONE] stealth font -> {output_dir}/fonts/0.woff + ttffonts/0.ttf")


def createfonts(reference_font, output_dir, font_name):
    """Generate one Evil Font variant per character in LETTERS.
    
    In each variant, every character's glyph is replaced with the glyph for
    `currentletter`. This means that no matter what Unicode byte is stored in
    the document, the renderer will draw `currentletter` — the core Evil Font trick.

    Also writes all @font-face rules to fonts.css (overwrites any existing file).
    """
    log(f"Source font:  {reference_font}")
    log(f"Output dir:   {output_dir}")
    log(f"Characters:   {len(LETTERS)} variants to generate")

    with open(f'{output_dir}/fonts.css', "w") as css_file:

        for currentletter in LETTERS:
            # Load a fresh copy of the font for each variant to avoid cross-contamination
            font = TTFont(reference_font)
            font.recalcBBoxes = False

            unicode_cmap = _get_unicode_cmap(font)

            # Get the source glyph and its advance width for this letter
            source_glyph_name = unicode_cmap.cmap.get(ord(currentletter))
            source_width = font['hmtx'].metrics[source_glyph_name][0]

            # Take a deep copy of the source glyph to use as a stamp
            source_glyph = copy.deepcopy(font['glyf'][source_glyph_name])

            # Remap every other character in LETTERS to look like currentletter
            for table in font['cmap'].tables:
                for char in LETTERS:
                    if char == currentletter:
                        continue
                    if ord(char) not in table.cmap:
                        continue

                    target_glyph_name = table.cmap[ord(char)]
                    target_original = font['glyf'][target_glyph_name]

                    # Preserve the target's original bounding box so that
                    # spacing and baseline positioning remain correct
                    orig_xMin = getattr(target_original, 'xMin', 0)
                    orig_yMax = getattr(target_original, 'yMax', 0)
                    orig_yMin = getattr(target_original, 'yMin', 0)

                    # Stamp a copy of the source glyph into the target slot
                    font['glyf'][target_glyph_name] = copy.deepcopy(source_glyph)

                    # Restore the vertical bounds (keeps line height consistent)
                    font['glyf'][target_glyph_name].yMax = orig_yMax
                    font['glyf'][target_glyph_name].yMin = orig_yMin

                    # Match advance width to source; preserve original LSB for positioning
                    font['hmtx'].metrics[target_glyph_name] = (source_width, orig_xMin)

            _remove_layout_tables(font)

            # Each font variant is named using the hex encoding of the letter
            # so the name is unique and safely usable as a filename
            letter_hex = currentletter.encode().hex()
            new_name = f'{font_name} {letter_hex}'
            _rename_font(font, new_name)

            # Save as WOFF for web use and TTF for document/desktop use
            font.flavor = 'woff'
            font.save(f'{output_dir}/fonts/{letter_hex}.woff')
            font.flavor = None
            font.save(f'{output_dir}/ttffonts/{letter_hex}.ttf')

            # Write the @font-face rule for this variant
            css_file.write(
                f'@font-face {{'
                f'font-family: "{letter_hex}";'
                f'src: url("fonts/{letter_hex}.woff") format(\'woff\');'
                f'}}'
            )

            log(f"  [{letter_hex}] '{currentletter}' -> {output_dir}/fonts/{letter_hex}.woff + ttffonts/{letter_hex}.ttf")


# ---------------------------------------------------------------------------
# HTML output
# ---------------------------------------------------------------------------

def _write_html(template_file, output_file, content):
    """Inject `content` into the HTML template and write to `output_file`."""
    with open(template_file, "r") as f:
        html = f.read()
    html = html.replace("<!-- #STUFF HERE -->", content)
    with open(output_file, "w") as f:
        f.write(html)


def createhtml(input_human_file, input_computer_file, output_file):
    """Build a steganographic HTML file from human and computer text files.
    
    Each character from the computer file is wrapped in a <span> that applies
    an Evil Font chosen by the corresponding human file character. To a human
    reading the rendered page, the text looks like the human file. A machine
    parsing the raw HTML sees the computer file.

    Lines where the computer text is shorter than the human text are skipped.
    Extra computer characters (beyond the human line length) are hidden using
    the stealth font (font-family: '0').
    """
    log(f"Human file:    {input_human_file}")
    log(f"Computer file: {input_computer_file}")
    log(f"Output file:   {output_file}")

    spans = []
    lines_processed = 0
    total_hidden = 0
    line_number = 0

    with open(input_human_file, "r") as human_file, \
         open(input_computer_file, "r") as computer_file:

        for human_line, computer_line in zip(
            (l.rstrip('\n') for l in human_file),
            (l.rstrip('\n') for l in computer_file),
        ):
            h_len = len(human_line)
            c_len = len(computer_line)

            line_number += 1
            log(f"Human   ({h_len} chars): {human_line}")
            log(f"Computer({c_len} chars): {computer_line}")

            if c_len < h_len:
                log(f"  Skipping line {line_number}: computer line must be >= human line length.", verbose_only=False)
                continue

            # Extra computer characters are inserted at the midpoint of the human text
            diff = c_len - h_len
            mid = h_len // 2

            log(f"  diff={diff}, hidden chars inserted at mid={mid}")

            line_spans = []
            h_index = 0

            for i, computer_char in enumerate(computer_line):
                if mid <= i < mid + diff:
                    # Hidden character — rendered invisibly using the stealth font
                    line_spans.append(
                        f"<span style=\"font-family: '0';\">{computer_char}</span>"
                    )
                else:
                    # Visible character — disguised as the corresponding human character
                    human_char = human_line[h_index]
                    letter_hex = human_char.encode().hex()
                    line_spans.append(
                        f"<span style=\"font-family: '{letter_hex}';\">{computer_char}</span>"
                    )
                    h_index += 1

            spans.append("".join(line_spans))
            lines_processed += 1
            total_hidden += diff

    _write_html("template.html", output_file, "\n<br>\n".join(spans))
    log(f"[DONE] HTML written -> {output_file} ({lines_processed} lines, {total_hidden} hidden chars)")


# ---------------------------------------------------------------------------
# DOCX output
# ---------------------------------------------------------------------------

def create_doc(input_human_file, input_computer_file, output_file, font_name_in):
    """Build a steganographic DOCX file from human and computer text files.
    
    Works identically to createhtml() but outputs a Word document. Each run is
    assigned an Evil Font by setting all four font slots (ascii, hAnsi, eastAsia,
    cs) to prevent Word from falling back to a system font.

    The TTF variants of the Evil Fonts must be installed on the system (or
    embedded in the document) for the deception to render correctly in Word.

    Lines where the computer text is shorter than the human text are skipped.
    """
    log(f"Human file:    {input_human_file}")
    log(f"Computer file: {input_computer_file}")
    log(f"Output file:   {output_file}")
    log(f"Font family:   {font_name_in}")

    doc = Document()
    lines_processed = 0
    total_hidden = 0
    line_number = 0

    with open(input_human_file, "r") as human_file, \
         open(input_computer_file, "r") as computer_file:

        for human_line, computer_line in zip(
            (l.rstrip('\n') for l in human_file),
            (l.rstrip('\n') for l in computer_file),
        ):
            h_len = len(human_line)
            c_len = len(computer_line)

            line_number += 1
            log(f"Human   ({h_len} chars): {human_line}")
            log(f"Computer({c_len} chars): {computer_line}")

            if c_len < h_len:
                log(f"  Skipping line {line_number}: computer line must be >= human line length.", verbose_only=False)
                continue

            diff = c_len - h_len
            mid = h_len // 2

            log(f"  diff={diff}, hidden chars inserted at mid={mid}")

            p = doc.add_paragraph()
            h_index = 0

            for i, computer_char in enumerate(computer_line):
                if mid <= i < mid + diff:
                    # Hidden character — use the stealth (zero-width) font
                    font_name = f'{font_name_in} 0'
                else:
                    # Visible character — disguised as the corresponding human character
                    human_char = human_line[h_index]
                    font_name = f'{font_name_in} {human_char.encode().hex()}'
                    h_index += 1

                # Add a run and explicitly set all four font slots.
                # Word will fall back to a system font if any slot is unset,
                # which would break the illusion.
                run = p.add_run(computer_char)
                run.font.name = font_name
                rFonts = run._element.rPr.rFonts
                rFonts.set(qn("w:ascii"),   font_name)
                rFonts.set(qn("w:hAnsi"),   font_name)
                rFonts.set(qn("w:eastAsia"), font_name)
                rFonts.set(qn("w:cs"),      font_name)

            lines_processed += 1
            total_hidden += diff

    doc.save(output_file)
    log(f"[DONE] DOCX written -> {output_file} ({lines_processed} lines, {total_hidden} hidden chars)")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def setup_parser():
    parser = argparse.ArgumentParser(
        description="Fontuscator — Evil Font steganography tool for security research.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument(
        '-v', '--verbose',
        action='store_true',
        help='Enable verbose output. Without this flag the tool runs silently.',
    )
    subparsers = parser.add_subparsers(dest='command', required=True)

    # --- create ---
    create_parser = subparsers.add_parser(
        'create',
        help='Generate an Evil Font family from a reference font.',
        description=(
            'Produces one font variant per printable character plus a stealth font, '
            'along with a fonts.css file for web use.'
        ),
    )
    create_parser.add_argument('reference_font', help='Path to the source .ttf font file.')
    create_parser.add_argument('output_dir',     help='Directory to write fonts and CSS into.')
    create_parser.add_argument('font_name',      help='Internal name prefix for the font family.')

    # --- web ---
    web_parser = subparsers.add_parser(
        'web',
        help='Generate a steganographic HTML file.',
        description=(
            'Wraps computer-file characters in Evil Font spans so the page displays '
            'human-file text visually while the raw HTML contains the computer text. '
            'fonts.css must be present in the output directory.'
        ),
    )
    web_parser.add_argument('input_human_file',    help='Text visible to human readers.')
    web_parser.add_argument('input_computer_file', help='Text visible to machines / AI.')
    web_parser.add_argument('output_file',         help='Path for the generated HTML file.')

    # --- doc ---
    doc_parser = subparsers.add_parser(
        'doc',
        help='Generate a steganographic DOCX file.',
        description=(
            'Produces a Word document using Evil Fonts so the displayed text differs '
            'from the underlying Unicode. font_name must match the value used in create.'
        ),
    )
    doc_parser.add_argument('input_human_file',    help='Text visible to human readers.')
    doc_parser.add_argument('input_computer_file', help='Text visible to machines / AI.')
    doc_parser.add_argument('output_file',         help='Path for the generated DOCX file.')
    doc_parser.add_argument('font_name',           help='Font family name (must match create step).')

    return parser


def main():
    parser = setup_parser()
    args = parser.parse_args()

    # Set the global verbose flag before any functions run
    global VERBOSE
    VERBOSE = args.verbose

    if args.command == 'create':
        if not os.path.isfile(args.reference_font):
            print(f"Error: '{args.reference_font}' does not exist.")
            sys.exit(1)

        # Create output directories if they don't already exist
        for subdir in ('', '/fonts', '/ttffonts'):
            path = args.output_dir + subdir
            if not os.path.isdir(path):
                log(f"Creating directory: {path}")
                os.makedirs(path, exist_ok=True)

        createfonts(args.reference_font, args.output_dir, args.font_name)
        createstealthfont(args.reference_font, args.output_dir, args.font_name)

    elif args.command == 'web':
        for path in (args.input_human_file, args.input_computer_file):
            if not os.path.isfile(path):
                print(f"Error: '{path}' does not exist.")
                sys.exit(1)
        createhtml(args.input_human_file, args.input_computer_file, args.output_file)

    elif args.command == 'doc':
        for path in (args.input_human_file, args.input_computer_file):
            if not os.path.isfile(path):
                print(f"Error: '{path}' does not exist.")
                sys.exit(1)
        create_doc(
            args.input_human_file,
            args.input_computer_file,
            args.output_file,
            args.font_name,
        )


if __name__ == "__main__":
    main()